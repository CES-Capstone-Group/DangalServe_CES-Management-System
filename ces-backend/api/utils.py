import os
import platform
from docx import Document
from django.conf import settings
from docx2pdf import convert

# Only import pythoncom if running on Windows
if platform.system() == "Windows":
    import pythoncom

def convert_docx_to_pdf(docx_path):
    # Check if the current system is Windows
    if platform.system() == "Windows":
        # Initialize COM objects using pythoncom for Windows
        pythoncom.CoInitialize()
        try:
            # Define the output path for the PDF
            pdf_output_path = docx_path.replace('.docx', '.pdf')

            # Perform the conversion
            convert(docx_path, pdf_output_path)

            return pdf_output_path
        finally:
            # Uninitialize COM objects
            pythoncom.CoUninitialize()
    else:
        # If not Windows, raise an error or handle the conversion differently
        raise NotImplementedError("DOCX to PDF conversion is only supported on Windows.")

def replace_text_in_paragraph(paragraph, placeholders):
    """Replaces placeholders in a given paragraph with actual values."""
    for key, value in placeholders.items():
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, value or '')

def replace_text_in_table(table, placeholders):
    """Replaces placeholders in all cells of a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, placeholders)

def generate_proposal_doc(proposal):
    # Path to your Word document template
    template_path = os.path.join(settings.BASE_DIR, 'Document Formats', 'PNC-PRE-FO-34-Community-and-Extension-Service-Activity-Proposal.docx')
    
    # Load the document template
    doc = Document(template_path)
    
    # Convert partner_community to a displayable format as a list
    partner_communities_display = ', '.join([community.strip() for community in proposal.partner_community.split(',')]) if proposal.partner_community else 'N/A'

    # Handle Identified Needs (Text or File)
    if proposal.identified_needs_file:
        identified_needs = f"See attached file: {os.path.basename(proposal.identified_needs_file.name)}"
    else:
        identified_needs = proposal.identified_needs_text or 'N/A'

    # Handle Budget Requirement (Text or File)
    if proposal.budget_requirement_file:
        budget_requirement = f"See attached file: {os.path.basename(proposal.budget_requirement_file.name)}"
    else:
        budget_requirement = proposal.budget_requirement_text or 'N/A'

    # Proponents (Display all proponents - name and position)
    proponent_details = "\n\n\n".join([f"{proponent.name} - {proponent.position}" for proponent in proposal.proponents.all()]) or 'N/A'

    # Signatories (Multiple signatories for each section)
    prepared_by_signatories = "\n\n\n".join([f"{signatory.name} - {signatory.position}" for signatory in proposal.signatories.filter(section='prepared')]) or '_____________'
    endorsed_by_signatories = "\n\n\n".join([f"{signatory.name} - {signatory.position}" for signatory in proposal.signatories.filter(section='endorsed')]) or '_____________'
    concurred_by_signatories = "\n\n\n".join([f"{signatory.name} - {signatory.position}" for signatory in proposal.signatories.filter(section='concurred')]) or '_____________'

    three_year_checkbox = '☑' if proposal.is_three_year_plan else '☐'
    one_year_checkbox = '☑' if proposal.is_one_year_plan else '☐'
    # Prepare placeholders and their replacements
    placeholders = {
        '{{three_year_checkbox}}': three_year_checkbox,  # Three-Year checkbox
        '{{one_less_checkbox}}': one_year_checkbox, # One-Year checkbox
        '{{title}}': proposal.title,
        '{{engagement_date}}': str(proposal.engagement_date) if proposal.engagement_date else '',
        '{{disengagement_date}}': str(proposal.disengagement_date) if proposal.disengagement_date else '',
        '{{department_organization}}': proposal.department,
        '{{lead_proponent}}': proponent_details,  # Display all proponents
        '{{contact_details}}': proposal.contact_details,
        '{{project_description}}': proposal.project_description,
        '{{target_date}}': str(proposal.target_date) if proposal.target_date else '',
        '{{location}}': proposal.location,
        '{{school}}': '☑' if proposal.school else '☐',
        '{{barangay}}': '☑' if proposal.barangay else '☐',
        '{{gov_org}}': f"{proposal.government_org}" if proposal.government_org else '_____________',
        '{{non_gov_org}}': f"{proposal.non_government_org}" if proposal.non_government_org else '_____________',
        '{{gov_org_box}}': '☑' if proposal.government_org else '☐',
        '{{non_gov_box}}': '☑' if proposal.non_government_org else '☐',
        '{{partner_community}}': partner_communities_display,
        '{{identified_needs}}': identified_needs,
        '{{budget_requirement}}': budget_requirement,
        '{{gen_objs}}': proposal.general_objectives,
        '{{spec_objs}}': proposal.specific_objectives,
        '{{success_indicators}}': proposal.success_indicators,
        '{{cooperating_agencies}}': proposal.cooperating_agencies,
        '{{monitoring_mech}}': proposal.monitoring_mechanics,
        '{{evaluation_mech}}': proposal.evaluation_mechanics,
        '{{timetable}}': proposal.timetable,
        '{{risk_assessment}}': proposal.risk_assessment,
        '{{action_plans_to_address_risks}}': proposal.action_plans,
        '{{sustainability_approaches}}': proposal.sustainability_approaches,
        '{{prepared_by}}': prepared_by_signatories,  # Prepared By multiple signatories
        '{{endorsed_by}}': endorsed_by_signatories,  # Endorsed By multiple signatories
        '{{concurred_by}}': concurred_by_signatories  # Concurred By multiple signatories
    }

    # Handle partner barangay signatures if applicable
    partner_signatures = ''
    for barangay_approval in proposal.barangay_approvals.filter(status='Approved'):
        partner_signatures += f"{barangay_approval.barangay_name} - Signed by on {barangay_approval.sign_date}\n\n\n"
    
    placeholders['{{partner_signatures}}'] = partner_signatures or '_____________'

    # Replace placeholders in all paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, placeholders)

    # Replace placeholders in tables
    for table in doc.tables:
        replace_text_in_table(table, placeholders)

    # Save the generated document
    output_dir = os.path.join(settings.MEDIA_ROOT, 'generated_proposals')
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, f'proposal_{proposal.proposal_id}.docx')
    doc.save(output_path)
    
    return output_path
